from pathlib import Path

from docx import Document

from cleantext_studio.cleaners import clean_text
from cleantext_studio.exporters import export_docx_blocks


def test_docx_reopens_with_heading_list_table_and_chinese(tmp_path: Path) -> None:
    result = clean_text("## 项目\n\n- 农户\n\n|模块|功能|\n|-|-|\n|识别|诊断|")
    path = tmp_path / "验收.docx"
    export_docx_blocks(result.blocks, path)
    reopened = Document(path)
    assert len(reopened.tables) == 1
    assert "项目" in "".join(p.text for p in reopened.paragraphs)
    assert "农户" in "".join(p.text for p in reopened.paragraphs)
