from zipfile import ZipFile

from docx import Document

from cleantext_studio.cleaners import clean_text
from cleantext_studio.exporters import export_docx_blocks
from cleantext_studio.models import TextBlockType

SOURCE = r'''### 社会熵变模型

我们可以用一个简化的反馈模型来描述社会秩序（低熵）与混乱（高熵）之间的动态平衡：

\[
\frac{dO}{dt}=I-\lambda O+\xi(t)
\]

这里，\( O \) 表示社会有序度（Order），\( I \) 表示制度和教育等带来的秩序输入，\( \lambda \) 是系统内在的衰退系数，而 \( \xi(t) \) 则是随机冲击。

| **社会熵类型** | **表现形式** | **负熵（有序化）手段** | **Emoji** |
| --- | --- | --- | --- |
| **认知熵** | 信息茧房、后真相、谣言肆虐 | 批判性思维教育、事实核查、开源情报 | 🧠➡️💡 |
| **结构熵** | 官僚主义、流程繁琐、行政低效 | 扁平化管理、去中心化自治组织、数字化流程 | 🏛️➡️⚡ |
'''


def test_social_entropy_docx_has_native_inline_math_and_clean_table(tmp_path) -> None:
    result = clean_text(SOURCE)
    heading = next(block for block in result.blocks if block.text == "社会熵变模型")
    assert heading.type == TextBlockType.HEADING_3
    table = next(block.table for block in result.blocks if block.table)
    assert table.headers == ["社会熵类型", "表现形式", "负熵(有序化)手段"]
    assert all("**" not in value for row in [table.headers, *table.rows] for value in row)

    path = tmp_path / "social-entropy.docx"
    export_docx_blocks(result.blocks, path)
    document = Document(path)
    assert len(document.tables) == 1
    assert len(document.tables[0].columns) == 3
    with ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert xml.count("<m:oMath") >= 5
    assert "<m:oMathPara" in xml
    assert all(item not in xml for item in (r"\(", r"\)", r"\lambda", r"\xi", "**"))
    assert "🧠" not in xml
