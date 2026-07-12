from pathlib import Path
from zipfile import ZipFile

from docx import Document

from cleantext_studio.cleaners import clean_text
from cleantext_studio.exporters import export_docx_blocks


def test_native_docx_table_structure(tmp_path: Path) -> None:
    result = clean_text("# 标题\n\n|模块|功能|\n|-|-|\n|识别|诊断|\n\n1. 有序项")
    path = tmp_path / "structure.docx"
    export_docx_blocks(result.blocks, path)
    doc = Document(path)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 2 and len(doc.tables[0].columns) == 2
    assert doc.tables[0].cell(1, 0).text == "识别"
    assert any(p.style.name.startswith("Heading") for p in doc.paragraphs)
    with ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    assert b"<w:tbl>" in xml
