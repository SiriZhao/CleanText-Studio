from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from cleantext_studio.cleaners import clean_text
from cleantext_studio.exporters import export_docx_blocks
from cleantext_studio.models import TextBlockType


@pytest.mark.parametrize(
    "source",
    [
        "| 项目 | 金额 |\n|---|---|\n|开发|8万元|",
        "|项目|金额|\n|-|-|\n|开发|8万元|",
        "|Year|Revenue|\n|---:|:---|\n|2026|30万元|",
        "|模块|功能|\n|-|-|\n|AI病虫害识别|拍照识别病害、虫害、草害并生成较长的辅助说明|",
    ],
)
def test_markdown_table_variants(source: str) -> None:
    result = clean_text(source)
    tables = [block for block in result.blocks if block.type == TextBlockType.TABLE]
    assert len(tables) == 1
    assert tables[0].table is not None
    assert tables[0].table.headers[0] in {"项目", "Year", "模块"}
    assert tables[0].table.rows


def test_empty_cells_and_mixed_markdown() -> None:
    source = "## 标题\n\n|项目|说明|\n|-|-|\n|开发||\n||空项目|\n\n- 列表"
    result = clean_text(source)
    table = next(block.table for block in result.blocks if block.table)
    assert table.rows == [["开发", ""], ["", "空项目"]]
    assert any(block.type == TextBlockType.HEADING_2 for block in result.blocks)
    assert any(block.type == TextBlockType.LIST_ITEM for block in result.blocks)


def test_docx_renderer_creates_real_table(tmp_path: Path) -> None:
    result = clean_text(
        "## 创业方案\n\n|年份|收入|利润|\n|-|-|-|\n|第一年|34万|-33万|\n\n- 保留列表"
    )
    path = tmp_path / "test_output.docx"
    export_docx_blocks(result.blocks, path)
    document = Document(path)
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "年份"
    assert any(p.style.name == "List Bullet" for p in document.paragraphs)
    with ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    assert b"<w:tbl>" in xml


def test_complete_business_plan_table() -> None:
    source = """# 县域农业创业方案

## 产品模块

| 模块 | 功能 |
| --- | --- |
| AI病虫害识别 | 图片识别 |
| 农事处方生成 | 用药建议 |

> 本项目坚持本地优先。

- 服务农户
- 服务合作社"""
    result = clean_text(source)
    assert len([b for b in result.blocks if b.table]) == 1
    assert "AI病虫害识别" in result.text
