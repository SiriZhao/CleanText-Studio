from __future__ import annotations

import os
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from cleantext_studio.cleaners.tables import TableWidthPlanner
from cleantext_studio.math import OMMLConverter
from cleantext_studio.models import (
    DocumentTemplate,
    InlineRun,
    InlineRunType,
    MathExportMode,
    TextBlock,
    TextBlockType,
)


@dataclass(slots=True, frozen=True)
class DocumentExportQuality:
    level: str
    native_formulas: int
    fallback_formulas: int
    tables: int
    markdown_residuals: int
    empty_columns_removed: int


def evaluate_export_quality(blocks: list[TextBlock], renderer: DocxRenderer) -> DocumentExportQuality:
    """Report real renderer outcomes instead of counting formula fallback as success."""
    tables = sum(block.table is not None for block in blocks)
    residuals = sum("markdown" in warning for block in blocks for warning in block.warnings)
    removed = sum(
        -item
        for block in blocks
        if block.table is not None
        for item in block.table.malformed_rows
        if item < 0
    )
    level = "优秀" if not renderer.formulas_as_text and not residuals else "良好"
    if renderer.formulas_as_text and renderer.formulas_as_omml == 0:
        level = "需要检查"
    return DocumentExportQuality(
        level,
        renderer.formulas_as_omml,
        renderer.formulas_as_text,
        tables,
        residuals,
        removed,
    )


def export_txt(text: str, path: Path, encoding: str = "utf-8", newline: str = "\r\n") -> None:
    data = text.replace("\r\n", "\n").replace("\r", "\n").replace("\n", newline)
    path.write_text(data, encoding=encoding, newline="")


def _font(run: object, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name  # type: ignore[attr-defined]
    run.font.size = Pt(size)  # type: ignore[attr-defined]
    run.font.bold = bold  # type: ignore[attr-defined]
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)  # type: ignore[attr-defined]


class DocxRenderer:
    """Render structured cleaning blocks into native Word elements."""

    def __init__(
        self,
        template: DocumentTemplate | None = None,
        math_export_mode: MathExportMode = MathExportMode.WORD_OMML,
    ) -> None:
        self.template = template or DocumentTemplate(
            title_size=22, body_size=12, line_spacing=1.5, first_line_chars=2, margin_cm=2.54
        )
        self.math_converter = OMMLConverter()
        self.math_export_mode = math_export_mode
        self.formulas_as_omml = 0
        self.formulas_as_text = 0

    def render(self, blocks: list[TextBlock]) -> DocumentType:
        doc = Document()
        section = doc.sections[0]
        margin = Cm(self.template.margin_cm)
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = (
            margin
        )
        for block in blocks:
            if not block.text:
                continue
            if block.type == TextBlockType.TABLE and block.table:
                self._table(doc, block)
                continue
            if block.math is not None:
                self._math_paragraph(doc, block)
                continue
            self._paragraph(doc, block)
        self._header_footer(doc)
        return doc

    def _paragraph(self, doc: DocumentType, block: TextBlock) -> None:
        heading_styles = {
            TextBlockType.TITLE: "Title",
            TextBlockType.HEADING_1: "Heading 1",
            TextBlockType.HEADING_2: "Heading 2",
            TextBlockType.HEADING_3: "Heading 3",
        }
        style = heading_styles.get(block.type)
        if block.type == TextBlockType.LIST_ITEM:
            style = "List Bullet"
        if block.type == TextBlockType.ORDERED_LIST_ITEM:
            style = "List Number"
        if block.type == TextBlockType.QUOTE:
            style = "Quote"
        paragraph = doc.add_paragraph(style=style)
        text = (
            block.text[2:]
            if block.type == TextBlockType.LIST_ITEM and block.text.startswith("• ")
            else block.text
        )
        runs = block.runs
        if block.type == TextBlockType.ORDERED_LIST_ITEM and block.list_marker:
            # Word owns ordered-list markers. Keeping the Markdown marker in
            # the run text creates duplicate labels such as "1. 1.".
            marker = block.list_marker
            if text.startswith(marker):
                text = text[len(marker) :].lstrip()
                runs = self._strip_run_prefix(runs, marker)
        if runs and any(item.type == InlineRunType.INLINE_MATH for item in runs):
            self._inline_runs(paragraph, runs)
            run = None
        else:
            run = paragraph.add_run(text)
        is_heading = block.type in heading_styles
        if run is not None:
            _font(
                run,
                self.template.title_font if is_heading else self.template.body_font,
                self.template.title_size if is_heading else self.template.body_size,
                is_heading,
            )
        paragraph.paragraph_format.line_spacing = self.template.line_spacing
        if is_heading:
            paragraph.paragraph_format.keep_with_next = True
        elif block.type == TextBlockType.PARAGRAPH:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Pt(
                self.template.body_size * self.template.first_line_chars
            )

    @staticmethod
    def _strip_run_prefix(runs: list[InlineRun], marker: str) -> list[InlineRun]:
        remaining = marker
        output: list[InlineRun] = []
        for item in runs:
            if remaining and item.type == InlineRunType.TEXT and item.text.startswith(remaining):
                text = item.text[len(remaining) :].lstrip()
                remaining = ""
                if text:
                    output.append(InlineRun(item.type, text, item.math))
            else:
                output.append(item)
        return output

    def _math_paragraph(self, doc: DocumentType, block: TextBlock) -> None:
        assert block.math is not None
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True
        expression = block.math.expression_source or block.math.normalized_text
        result = self.math_converter.convert(expression, display=True)
        if self.math_export_mode != MathExportMode.WORD_OMML:
            result.converted = False
        if result.converted and result.element is not None:
            paragraph._p.append(result.element)
            self.formulas_as_omml += 1
        else:
            run = paragraph.add_run(result.fallback_text)
            _font(run, "Cambria Math", self.template.body_size)
            self.formulas_as_text += 1
        if block.math.equation_number:
            paragraph.add_run(f"  {block.math.equation_number}")

    def _inline_runs(self, paragraph: object, runs: list[InlineRun]) -> None:
        for item in runs:
            if item.type != InlineRunType.INLINE_MATH:
                run = paragraph.add_run(item.text)  # type: ignore[attr-defined]
                _font(run, self.template.body_font, self.template.body_size)
                continue
            formula = item.math
            assert formula is not None
            result = self.math_converter.convert(formula.expression_source or formula.normalized_text)
            if self.math_export_mode != MathExportMode.WORD_OMML:
                result.converted = False
            if result.converted and result.element is not None:
                paragraph._p.append(result.element)  # type: ignore[attr-defined]
                self.formulas_as_omml += 1
                formula.conversion_status = "native_omml"
            else:
                run = paragraph.add_run(result.fallback_text)  # type: ignore[attr-defined]
                _font(run, "Cambria Math", self.template.body_size)
                self.formulas_as_text += 1
                formula.conversion_status = "unicode_fallback"

    def _table(self, doc: DocumentType, block: TextBlock) -> None:
        assert block.table is not None
        data = block.table
        table = doc.add_table(rows=1, cols=len(data.headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for column, value in enumerate(data.headers):
            self._cell(table.rows[0].cells[column], value, True)
        for row_data in data.rows:
            cells = table.add_row().cells
            for column, value in enumerate(row_data):
                self._cell(cells[column], value, False)
        section = doc.sections[0]
        page_width = section.page_width or Cm(21)
        left_margin = section.left_margin or Cm(self.template.margin_cm)
        right_margin = section.right_margin or Cm(self.template.margin_cm)
        available_width = page_width - left_margin - right_margin
        for column, ratio in enumerate(TableWidthPlanner().proportions(data)):
            width = available_width * ratio
            for row in table.rows:
                row.cells[column].width = width
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    def _cell(self, cell: object, text: str, bold: bool) -> None:
        cell.text = ""  # type: ignore[attr-defined]
        paragraph = cell.paragraphs[0]  # type: ignore[attr-defined]
        from cleantext_studio.math import MathDetector

        formulas = MathDetector().detect_inline(text)
        if formulas:
            cursor = 0
            for formula in formulas:
                if formula.start > cursor:
                    run = paragraph.add_run(text[cursor : formula.start])
                    _font(run, self.template.body_font, self.template.body_size, bold)
                result = self.math_converter.convert(formula.content)
                if self.math_export_mode != MathExportMode.WORD_OMML:
                    result.converted = False
                if result.converted and result.element is not None:
                    paragraph._p.append(result.element)
                    self.formulas_as_omml += 1
                else:
                    run = paragraph.add_run(result.fallback_text)
                    _font(run, "Cambria Math", self.template.body_size, bold)
                    self.formulas_as_text += 1
                cursor = formula.end
            if cursor < len(text):
                run = paragraph.add_run(text[cursor:])
                _font(run, self.template.body_font, self.template.body_size, bold)
        else:
            run = paragraph.add_run(text)
            _font(run, self.template.body_font, self.template.body_size, bold)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # type: ignore[attr-defined]

    def _header_footer(self, doc: DocumentType) -> None:
        section = doc.sections[0]
        if self.template.header:
            run = section.header.paragraphs[0].add_run(self.template.header)
            _font(run, self.template.body_font, 9)
        if self.template.page_numbers:
            paragraph = section.footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            field = OxmlElement("w:fldSimple")
            field.set(qn("w:instr"), "PAGE")
            paragraph._p.append(field)


def _atomic_save(document: DocumentType, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, name = tempfile.mkstemp(suffix=".docx", dir=path.parent)
    os.close(fd)
    temporary = Path(name)
    try:
        document.save(str(temporary))
        os.replace(temporary, path)
    finally:
        temporary.unlink(missing_ok=True)


def export_docx_blocks(
    blocks: list[TextBlock],
    path: Path,
    template: DocumentTemplate | None = None,
    math_export_mode: MathExportMode = MathExportMode.WORD_OMML,
) -> None:
    _atomic_save(DocxRenderer(template, math_export_mode).render(blocks), path)


def export_docx(text: str, path: Path, template: DocumentTemplate | None = None) -> None:
    from cleantext_studio.cleaners import clean_text

    export_docx_blocks(clean_text(text).blocks, path, template)
