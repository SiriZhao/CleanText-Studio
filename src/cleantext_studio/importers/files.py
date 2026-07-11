from __future__ import annotations

from pathlib import Path

from charset_normalizer import from_bytes
from docx import Document


def import_file(path: Path, max_mb: int = 20) -> str:
    if path.stat().st_size > max_mb * 1024 * 1024:
        raise ValueError(f"文件超过 {max_mb} MB 限制")
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        match = from_bytes(path.read_bytes()).best()
        if match is None:
            raise UnicodeError("无法识别文本编码")
        return str(match)
    if suffix == ".docx":
        doc = Document(str(path))
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("| " + " | ".join(cell.text for cell in row.cells) + " |")
        return "\n".join(parts)
    raise ValueError(f"不支持的文件类型：{suffix}")
